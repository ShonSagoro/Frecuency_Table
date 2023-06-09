from docx import Document
from docx.shared import Inches
from utilities import frecuency_table_tools as ftt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import csv
import os

csv_file=os.path.join(os.path.dirname(__file__), "csv_info/data.csv")
data = []

with open(csv_file) as archivo_csv:
    lector_csv = csv.reader(archivo_csv, delimiter=',')
    data=[float(number) for number in next(lector_csv)]
        
document=Document()

def write_a_frecuency_table(data_table:list, document:Document):
    table = document.add_table(rows=1, cols=10)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Clase'
    hdr_cells[1].text = 'Lim I'
    hdr_cells[2].text = 'Lim. S'
    hdr_cells[3].text = 'MC'
    hdr_cells[4].text = 'Frec'
    hdr_cells[5].text = 'Frec AC'
    hdr_cells[6].text = 'Lim IE'
    hdr_cells[7].text = 'Lim SE'
    hdr_cells[8].text = 'Frec R'
    hdr_cells[9].text = 'Frec R%'
    for data in data_table:
        id=0
        row_cells = table.add_row().cells
        for j in data:
            paragraph_data=row_cells[id].add_paragraph()
            paragraph_data_run=paragraph_data.add_run(str(j))
            paragraph_data_run.font.name = 'Century Gothic'
            id+=1
    document.add_paragraph("")


def write_a_table(info:str, data:list, document:Document):

    document.add_paragraph().add_run(info).bold=True
    document.add_paragraph(str(data))
    document.add_paragraph("")
    

def write_a_data(data, type:str, operation:str, document:Document):
    p = document.add_paragraph()
    prun= p.add_run(type+": ")
    prun.bold=True
    prun.font.name = 'Century Gothic'
    prun= p.add_run(str(operation))
    prun= p.add_run(str(data))

def add_image(imgName:str):
    p = document.add_paragraph()
    r = p.add_run()
    r.add_picture('public/'+imgName+'.png',width=Inches(5.00))
    last_element = document.paragraphs[-1]
    last_element.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Table with the number without sort
write_a_table("Not tidy.",data, document)

# Table sort
data.sort()
write_a_table("Tidy.", data, document)

# Add the range  
range=ftt.get_range(data)
write_a_data(range, "Range", str(data[len(data) - 1]) +"-"+ str(data[0])+"= ", document)

# Add number of classes
klases=ftt.get_klases(len(data))
write_a_data(klases, "Classes", "1 + 3.322 x log("+str(len(data))+")= ", document)

# Add amplitude
amplitude=ftt.get_amplitude(range, klases)
write_a_data(amplitude, "Amplitude" ,str(range)+"/"+str(klases)+"= ", document)

# Add variation unit
variation_unit=ftt.get_variation_unit(data)
write_a_data(variation_unit, "Variation Unit","", document)

# Make the frecuency table
frecuency_table=ftt.make_frecuency_table(data, klases, amplitude, variation_unit)

# Add Frecuency table
write_a_frecuency_table(frecuency_table, document)

media_arithmetic=ftt.get_arithmetic_media(frecuency_table, data)
write_a_data(media_arithmetic, "Media Arithmetic","", document)

moda=ftt.get_moda(frecuency_table, amplitude)
write_a_data(moda, "Moda","", document)

media=ftt.get_mediana(frecuency_table, data, amplitude)
write_a_data(media, "Media","", document)

klases=[]
frecuency=[]

for element in frecuency_table:
    klases.append(str(element[1])+"-"+str(element[2]))
    frecuency.append(element[3])


plt.bar(klases, frecuency)
plt.title("Tabla de Frecuencias")
plt.xlabel("Valores")
plt.ylabel("Frecuencias")

plt.savefig('public/frec_bar.png')
add_image('frec_bar')
plt.clf()


plt.pie(frecuency, labels=klases)
plt.title("Tabla de Frecuencias")

plt.savefig('public/frec_pie.png')
add_image('frec_pie')

write_a_data("Baljeet, Rusty, ShonSagoro", "Made by","", document)


document.save('public/frecuency_table.docx')
print("Document gerenerated successfully.")